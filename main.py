import os
from docx import Document
from instruction_parser import parse_contract_instructions
from docx_parser import extract_list_with_numbers, get_ilvl
import json

from paragraph_modifier import (
    add_numbered_paragraph,
    add_sentence_in_paragraph,
    normalize_paragraph
)

# === Directories ===
ORIGIN_DIR = "contracts"
DIST_DIR   = "updated"

# === Load and Parse Instruction Document ===
instruction_docx = Document(os.path.join(ORIGIN_DIR, 'Instructions and Snippets of Text.docx'))
instruction_text = "\n"
for para in instruction_docx.paragraphs:
    instruction_text += para.text + '\n'
instruction_data = parse_contract_instructions(instruction_text)

print(json.dumps(instruction_data, indent=2))

# === Process Each Instruction ===
for instruction in instruction_data:
    insert_type = instruction['insert_type']
    insert_text = instruction['text']
    section_name = instruction['section_name']
    contract_file = instruction['contract'] + '.docx'

    # Load document and normalize paragraph formatting
    doc_path = os.path.join(ORIGIN_DIR, contract_file)
    document = Document(doc_path)
    document = normalize_paragraph(document)

    # Extract numbered sections for indexing
    sections = extract_list_with_numbers(document)

    try:
        section_index = sections.index(section_name)

        if insert_type == 'as':
            # Insert a new numbered clause before the matched section
            ilvl = get_ilvl(section_name)
            document = add_numbered_paragraph(document, insert_text, index=section_index, num_id=1, ilvl=ilvl)

        elif insert_type == 'in':
            # Insert a sentence inside the matched section after sentence X
            sentence_after = instruction.get('sentence_after', 1)
            document = add_sentence_in_paragraph(document, insert_text, index=section_index, sentence_after=sentence_after)

    except ValueError:
        print(f"Unable to locate section '{section_name}' in {contract_file}. Skipping...")

    # Save updated document
    output_path = os.path.join(DIST_DIR, contract_file)
    document.save(output_path)
