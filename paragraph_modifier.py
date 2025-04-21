import re
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# === Utility: Copy formatting between paragraphs ===
def copy_paragraph_format(source, target):
    """Copy paragraph-level formatting (indentation, spacing, alignment)."""
    pf = source.paragraph_format
    tf = target.paragraph_format
    tf.left_indent = pf.left_indent
    tf.right_indent = pf.right_indent
    tf.first_line_indent = pf.first_line_indent
    tf.space_before = pf.space_before
    tf.space_after = pf.space_after
    tf.line_spacing = pf.line_spacing
    tf.alignment = pf.alignment

# === Utility: Copy font and styling from one run to another ===
def copy_run_format(source_run, target_run):
    """Copy font properties and styles from one run to another."""
    target_run.italic = source_run.italic
    target_run.underline = source_run.underline
    target_run.font.name = source_run.font.name
    target_run.font.size = source_run.font.size
    target_run.font.color.rgb = source_run.font.color.rgb

# === Add a numbered paragraph before a given index, apply formatting and bold text inside quotes ===
def add_numbered_paragraph(doc, text, index=0, num_id=1, ilvl=0):
    """Insert a numbered paragraph before a given index.
    - Preserves styles.
    - Bolds content inside smart quotes.
    - If base paragraph's first run is bold, first sentence is bold + underline (excluding the period).
    """

    base_paragraph = doc.paragraphs[index]
    new_paragraph = base_paragraph.insert_paragraph_before("", base_paragraph.style)

    copy_paragraph_format(base_paragraph, new_paragraph)

    # Apply numbering
    p_pr = new_paragraph._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')

    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), str(ilvl))
    numPr.append(ilvl_el)

    numId_el = OxmlElement('w:numId')
    numId_el.set(qn('w:val'), str(num_id))
    numPr.append(numId_el)

    p_pr.append(numPr)

    is_base_bold = base_paragraph.runs and base_paragraph.runs[0].bold
    base_run = base_paragraph.runs[0] if base_paragraph.runs else None

    # Split text by smart-quote-aware logic
    parts = re.split(r'(“[^”]+”)', text)
    sentence_chunks = []
    for part in parts:
        if part.startswith("“") and part.endswith("”"):
            sentence_chunks.append(part)
        else:
            sentence_chunks.extend(re.split(r'(?<=\.)\s+', part))

    for i, chunk in enumerate(filter(None, map(str.strip, sentence_chunks))):
        if chunk.startswith("“") and chunk.endswith("”"):
            run_open = new_paragraph.add_run("“")
            run_bold = new_paragraph.add_run(chunk[1:-1])
            run_bold.bold = True
            run_close = new_paragraph.add_run("” ")

            if base_run:
                copy_run_format(base_run, run_open)
                copy_run_format(base_run, run_bold)
                copy_run_format(base_run, run_close)

        elif i == 0 and is_base_bold:
            # Bold + underline only up to the first period
            period_index = chunk.find('.')
            if period_index != -1:
                # Underline up to before the period
                main_text = chunk[:period_index]
                dot = chunk[period_index]

                run_main = new_paragraph.add_run(main_text)
                if base_run:
                    copy_run_format(base_run, run_main)
                run_main.bold = True
                run_main.underline = True

                run_dot = new_paragraph.add_run(dot + " ")
                if base_run:
                    copy_run_format(base_run, run_dot)
                run_dot.bold = True
                run_dot.underline = False
            else:
                # No period → underline whole chunk
                run_all = new_paragraph.add_run(chunk + " ")
                if base_run:
                    copy_run_format(base_run, run_all)
                run_all.bold = True
                run_all.underline = True

        else:
            run = new_paragraph.add_run(chunk + " ")
            if base_run:
                copy_run_format(base_run, run)

    return doc

def find_nth_index(text, char, n):
    """Return the index of the nth occurrence of char in text."""
    index = -1
    for _ in range(n + 1):
        index = text.find(char, index + 1)
        if index == -1:
            return -1  # char not found enough times
    return index

# === Insert a sentence into a paragraph after a specific sentence number ===
def add_sentence_in_paragraph(doc, text, index=0, sentence_after=0):
    """Insert a sentence after a given sentence number in a paragraph."""
    para = doc.paragraphs[index]
    limit = 1
    for run in para.runs:
        sentences = run.text.split('.')
        #if the run has no sentences
        if len(sentences) == 1:
            continue
        #Run has sentences
        if len(sentences) > limit + sentence_after - 1:
            index = find_nth_index(text, '.', sentence_after - limit)
            run.text = run.text[:index] + ' ' + text + run.text[index:]
            break
        else:
            limit = len(sentences)

    return doc

# === Normalize numbered paragraph (e.g., "1. Text") and apply proper numId + ilvl ===
def normalize_paragraph(doc):
    """Detect paragraphs starting with a number (e.g., '1.') and convert them to real numbered paragraphs."""
    pattern = r'^(\d+[A-Za-z0-9\.]*)\.\s*(.*)'

    for p in doc.paragraphs:
        match = re.match(pattern, p.text)
        if match:
            # Clear original run content
            p.runs[0].clear()

            # Apply numbering info to paragraph XML
            p_pr = p._p.get_or_add_pPr()
            numPr = OxmlElement('w:numPr')

            ilvl_el = OxmlElement('w:ilvl')
            ilvl_el.set(qn('w:val'), str(0))
            numPr.append(ilvl_el)

            numId_el = OxmlElement('w:numId')
            numId_el.set(qn('w:val'), str(1))
            numPr.append(numId_el)

            p_pr.append(numPr)

    return doc
